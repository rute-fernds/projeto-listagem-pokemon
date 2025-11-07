from docx_parser_converter.docx_to_txt.docx_to_txt_converter import DocxToTxtConverter
from docx_parser_converter.docx_parsers.utils import read_binary_from_file_path
from docx import Document
import html

# ---------------- FUNÇÕES DE CONVERSÃO ----------------
def docx_to_txt(docx_path, txt_output_path):
    """ Converte um arquivo .docx para .txt """

    try:
        print(f"Lendo arquivo {docx_path}")
        docx_file_content = read_binary_from_file_path(docx_path)
        converter =  DocxToTxtConverter(docx_file_content, use_default_values=True)
        
        if converter:
            txt_output = converter.convert_to_txt(indent=True)
            converter.save_txt_to_file(txt_output, txt_output_path)
    except Exception as e:
        print(f"Não foi possível converter o arquivo {docx_path}\n Erro: {e}")

def txt_to_docx(txt_file, docx_output):
    """ Converte um arquivo .txt para .docx """
    try:
        document = Document()
        with open(txt_file, 'r', encoding='utf-8') as file:
            for line in file:
                paragraph_text = line.strip()

                if paragraph_text:
                    document.add_paragraph(paragraph_text)
                else:
                    document.add_paragraph('')   

        document.save(docx_output)
    except Exception as e:
        print(f"ERRO: {e}")

def txt_to_html(txt_file, html_output):
    """ Converte um arquivo .txt para .html básico """
    try:
        html_content = []

        # Esqueleto do .html
        html_content.append("<!DOCTYPE html>")
        html_content.append('<html lang="pt-br">')
        html_content.append("<head>")
        html_content.append('    <meta charset="UTF-8">')
        html_content.append('    <meta name="viewport" content="width=device-width, initial-scale=1.0">')
        html_content.append(f"    <title>Relatório de Análise</title>")
        
        # Adiciona estilo
        html_content.append("    <style>")
        html_content.append("        body { font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 20px auto; padding: 15px; border: 1px solid #ddd; border-radius: 5px; }")
        html_content.append("        h1 { color: #333; }")
        html_content.append("        p { margin: 10px 0; }")
        html_content.append("    </style>")
        
        html_content.append("</head>")
        html_content.append("<body>")

        # Leitura e conversão do .txt
        with open(txt_file, 'r', encoding='utf-8') as file:
            for line in file:
                paragraph_text = line.strip()

                if paragraph_text:
                    safe_text = html.escape(paragraph_text)
                    html_content.append(f"    <p>{safe_text}</p>")
                else:

                    html_content.append("    <p>&nbsp;</p>")

        # Fecha as tags do html
        html_content.append("</body>")
        html_content.append("</html>")

        # Salva o arquivo final
        with open(html_output, 'w', encoding='utf-8') as out_file:
            out_file.write("\n".join(html_content))
            
        print(f"Relatório HTML salvo com sucesso em {html_output}")

    except Exception as e:
        print(f"ERRO ao converter .txt para .html: {e}")

